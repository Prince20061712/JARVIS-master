"""
Advanced syllabus parser with multi-format support and intelligent structure extraction
"""

import re
import json
import yaml
from typing import List, Dict, Any, Optional, Tuple, Set, Union
from dataclasses import dataclass, field
from datetime import datetime, timedelta
from pathlib import Path
import logging
import hashlib
from collections import defaultdict
import PyPDF2
from bs4 import BeautifulSoup
import markdown
import nltk
from nltk.tokenize import sent_tokenize
import networkx as nx
from enum import Enum

# Download required NLTK data
try:
    nltk.data.find('tokenizers/punkt')
except LookupError:
    nltk.download('punkt')

class SyllabusFormat(Enum):
    """Supported syllabus formats"""
    PDF = "pdf"
    JSON = "json"
    YAML = "yaml"
    TEXT = "text"
    HTML = "html"
    MD = "markdown"
    DOCX = "docx"
    UNKNOWN = "unknown"

@dataclass
class Topic:
    """Represents a syllabus topic with metadata"""
    name: str
    level: int = 0
    parent: Optional[str] = None
    children: List[str] = field(default_factory=list)
    prerequisites: List[str] = field(default_factory=list)
    learning_objectives: List[str] = field(default_factory=list)
    key_concepts: List[str] = field(default_factory=list)
    estimated_hours: float = 1.0
    difficulty: float = 0.5
    importance: float = 0.5
    keywords: Set[str] = field(default_factory=set)
    references: List[str] = field(default_factory=list)
    metadata: Dict[str, Any] = field(default_factory=dict)
    
    def __hash__(self):
        return hash(self.name)

@dataclass
class Syllabus:
    """Complete syllabus structure"""
    title: str
    subject: str
    grade_level: str
    board: str
    topics: Dict[str, Topic] = field(default_factory=dict)
    topic_hierarchy: Dict[str, List[str]] = field(default_factory=dict)
    prerequisites: Dict[str, List[str]] = field(default_factory=dict)
    learning_outcomes: List[str] = field(default_factory=list)
    assessment_structure: Dict[str, Any] = field(default_factory=dict)
    textbooks: List[str] = field(default_factory=list)
    reference_materials: List[str] = field(default_factory=list)
    total_hours: float = 0.0
    exam_pattern: Dict[str, Any] = field(default_factory=dict)
    metadata: Dict[str, Any] = field(default_factory=dict)
    
    def __post_init__(self):
        self.total_hours = sum(t.estimated_hours for t in self.topics.values())

class SyllabusParser:
    """
    Advanced syllabus parser with intelligent structure detection and topic extraction
    """
    
    def __init__(self, config: Optional[Dict[str, Any]] = None):
        """
        Initialize syllabus parser
        
        Args:
            config: Configuration dictionary
        """
        self.config = config or {}
        
        # Pattern compilation
        self.topic_patterns = self._compile_topic_patterns()
        self.objective_patterns = self._compile_objective_patterns()
        self.assessment_patterns = self._compile_assessment_patterns()
        
        # Common educational boards
        self.boards = {
            'cbse': 'Central Board of Secondary Education',
            'icse': 'Indian Certificate of Secondary Education',
            'igcse': 'International General Certificate of Secondary Education',
            'ib': 'International Baccalaureate',
            'state': 'State Board',
            'gcse': 'General Certificate of Secondary Education',
            'a-level': 'Advanced Level',
            'ap': 'Advanced Placement'
        }
        
        # Grade level patterns
        self.grade_patterns = {
            r'class\s+(\d+)': r'Class \1',
            r'grade\s+(\d+)': r'Grade \1',
            r'year\s+(\d+)': r'Year \1',
            r'standard\s+(\d+)': r'Standard \1'
        }
        
        self.logger = logging.getLogger(__name__)
        
    async def parse(
        self,
        source: Union[str, Path, Dict],
        format: Optional[SyllabusFormat] = None
    ) -> Syllabus:
        """
        Parse syllabus from various sources
        
        Args:
            source: File path, URL, or dictionary
            format: Syllabus format (auto-detected if not provided)
            
        Returns:
            Syllabus object
        """
        # Detect format if not provided
        if not format:
            format = self._detect_format(source)
        
        # Parse based on format
        if format == SyllabusFormat.PDF:
            syllabus_data = await self._parse_pdf(source)
        elif format == SyllabusFormat.JSON:
            syllabus_data = await self._parse_json(source)
        elif format == SyllabusFormat.YAML:
            syllabus_data = await self._parse_yaml(source)
        elif format == SyllabusFormat.HTML:
            syllabus_data = await self._parse_html(source)
        elif format == SyllabusFormat.MD:
            syllabus_data = await self._parse_markdown(source)
        elif format == SyllabusFormat.DOCX:
            syllabus_data = await self._parse_docx(source)
        else:
            syllabus_data = await self._parse_text(source)
        
        # Extract structured information
        syllabus = await self._extract_syllabus_structure(syllabus_data)
        
        # Validate and enhance
        syllabus = await self._enhance_syllabus(syllabus)
        
        return syllabus
    
    def _detect_format(self, source: Union[str, Path, Dict]) -> SyllabusFormat:
        """Detect syllabus format from source"""
        if isinstance(source, dict):
            return SyllabusFormat.JSON
        
        source_str = str(source).lower()
        
        if source_str.startswith(('http://', 'https://')):
            return SyllabusFormat.HTML
        
        path = Path(source_str)
        if path.exists():
            suffix = path.suffix.lower()
            format_map = {
                '.pdf': SyllabusFormat.PDF,
                '.json': SyllabusFormat.JSON,
                '.yaml': SyllabusFormat.YAML,
                '.yml': SyllabusFormat.YAML,
                '.html': SyllabusFormat.HTML,
                '.htm': SyllabusFormat.HTML,
                '.md': SyllabusFormat.MD,
                '.markdown': SyllabusFormat.MD,
                '.txt': SyllabusFormat.TEXT,
                '.docx': SyllabusFormat.DOCX
            }
            return format_map.get(suffix, SyllabusFormat.TEXT)
        
        return SyllabusFormat.TEXT
    
    async def _parse_pdf(self, path: Union[str, Path]) -> Dict[str, Any]:
        """Parse PDF syllabus"""
        text = ""
        metadata = {}
        
        try:
            with open(path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Extract metadata
                if pdf_reader.metadata:
                    for key, value in pdf_reader.metadata.items():
                        metadata[key[1:]] = value
                
                # Extract text from each page
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                
        except Exception as e:
            self.logger.error(f"Error parsing PDF: {e}")
        
        return {
            'raw_text': text,
            'metadata': metadata,
            'num_pages': len(pdf_reader.pages)
        }
    
    async def _parse_json(self, source: Union[str, Path, Dict]) -> Dict[str, Any]:
        """Parse JSON syllabus"""
        if isinstance(source, dict):
            data = source
        else:
            with open(source, 'r', encoding='utf-8') as file:
                data = json.load(file)
        
        return data
    
    async def _parse_yaml(self, source: Union[str, Path]) -> Dict[str, Any]:
        """Parse YAML syllabus"""
        with open(source, 'r', encoding='utf-8') as file:
            data = yaml.safe_load(file)
        
        return data
    
    async def _parse_html(self, source: Union[str, Path]) -> Dict[str, Any]:
        """Parse HTML syllabus"""
        if str(source).startswith(('http://', 'https://')):
            import aiohttp
            async with aiohttp.ClientSession() as session:
                async with session.get(source) as response:
                    html = await response.text()
        else:
            with open(source, 'r', encoding='utf-8') as file:
                html = file.read()
        
        soup = BeautifulSoup(html, 'html.parser')
        
        # Remove script and style elements
        for script in soup(["script", "style"]):
            script.decompose()
        
        text = soup.get_text()
        
        return {
            'raw_text': text,
            'html': html,
            'title': soup.title.string if soup.title else None
        }
    
    async def _parse_markdown(self, source: Union[str, Path]) -> Dict[str, Any]:
        """Parse Markdown syllabus"""
        with open(source, 'r', encoding='utf-8') as file:
            md_text = file.read()
        
        # Convert to HTML for structure detection
        html = markdown.markdown(md_text)
        soup = BeautifulSoup(html, 'html.parser')
        text = soup.get_text()
        
        return {
            'raw_text': text,
            'markdown': md_text,
            'html': html
        }
    
    async def _parse_docx(self, source: Union[str, Path]) -> Dict[str, Any]:
        """Parse DOCX syllabus"""
        from docx import Document
        
        doc = Document(source)
        text = []
        
        for para in doc.paragraphs:
            if para.text:
                text.append(para.text)
        
        return {
            'raw_text': '\n'.join(text),
            'paragraphs': [p.text for p in doc.paragraphs if p.text]
        }
    
    async def _parse_text(self, source: Union[str, Path]) -> Dict[str, Any]:
        """Parse plain text syllabus"""
        if isinstance(source, str) and len(source) > 1000:  # Assume it's text content
            text = source
        else:
            with open(source, 'r', encoding='utf-8') as file:
                text = file.read()
        
        return {'raw_text': text}
    
    async def _extract_syllabus_structure(self, data: Dict[str, Any]) -> Syllabus:
        """Extract structured syllabus information"""
        
        # Get raw text
        if 'raw_text' in data:
            text = data['raw_text']
        else:
            text = json.dumps(data)
        
        # Extract basic metadata
        title = self._extract_title(text, data)
        subject = self._extract_subject(text, data)
        grade_level = self._extract_grade_level(text, data)
        board = self._extract_board(text, data)
        
        # Extract topics hierarchy
        topics, hierarchy = await self._extract_topics(text, data)
        
        # Extract prerequisites
        prerequisites = await self._extract_prerequisites(text, topics)
        
        # Extract learning outcomes
        learning_outcomes = self._extract_learning_outcomes(text)
        
        # Extract assessment structure
        assessment_structure = self._extract_assessment_structure(text)
        
        # Extract textbooks and references
        textbooks, references = self._extract_references(text)
        
        # Create syllabus object
        syllabus = Syllabus(
            title=title,
            subject=subject,
            grade_level=grade_level,
            board=board,
            topics=topics,
            topic_hierarchy=hierarchy,
            prerequisites=prerequisites,
            learning_outcomes=learning_outcomes,
            assessment_structure=assessment_structure,
            textbooks=textbooks,
            reference_materials=references,
            exam_pattern=self._extract_exam_pattern(text),
            metadata={
                'source_format': data.get('format', 'unknown'),
                'extracted_at': datetime.now().isoformat()
            }
        )
        
        return syllabus
    
    async def _extract_topics(
        self,
        text: str,
        data: Dict[str, Any]
    ) -> Tuple[Dict[str, Topic], Dict[str, List[str]]]:
        """Extract topics and build hierarchy"""
        topics = {}
        hierarchy = defaultdict(list)
        
        # Try different extraction strategies
        strategies = [
            self._extract_topics_from_json_structure,
            self._extract_topics_from_headers,
            self._extract_topics_from_bullets,
            self._extract_topics_from_numbering,
            self._extract_topics_from_toc
        ]
        
        for strategy in strategies:
            extracted = await strategy(text, data)
            if extracted:
                topics.update(extracted)
                break
        
        # If no topics found, fallback to sentence-based extraction
        if not topics:
            topics = await self._extract_topics_from_sentences(text)
        
        # Build hierarchy
        for topic_name, topic in topics.items():
            if topic.parent:
                hierarchy[topic.parent].append(topic_name)
        
        # Add root topics
        root_topics = [t for t in topics.values() if not t.parent]
        for topic in root_topics:
            hierarchy['root'].append(topic.name)
        
        return topics, dict(hierarchy)
    
    async def _extract_topics_from_json_structure(
        self,
        text: str,
        data: Dict[str, Any]
    ) -> Dict[str, Topic]:
        """Extract topics from JSON structure"""
        topics = {}
        
        def extract_from_dict(d: Dict, parent: Optional[str] = None, level: int = 0):
            for key, value in d.items():
                if isinstance(value, dict):
                    topic = Topic(
                        name=key,
                        level=level,
                        parent=parent
                    )
                    topics[key] = topic
                    extract_from_dict(value, key, level + 1)
                elif isinstance(value, list):
                    if all(isinstance(item, str) for item in value):
                        # List of subtopics
                        for item in value:
                            subtopic = Topic(
                                name=item,
                                level=level + 1,
                                parent=key
                            )
                            topics[item] = subtopic
                    elif all(isinstance(item, dict) for item in value):
                        # List of subtopic objects
                        for item in value:
                            if 'name' in item:
                                subtopic = Topic(
                                    name=item['name'],
                                    level=level + 1,
                                    parent=key,
                                    estimated_hours=item.get('hours', 1.0),
                                    difficulty=item.get('difficulty', 0.5)
                                )
                                topics[item['name']] = subtopic
        
        if isinstance(data, dict):
            extract_from_dict(data)
        
        return topics
    
    async def _extract_topics_from_headers(
        self,
        text: str,
        data: Dict[str, Any]
    ) -> Dict[str, Topic]:
        """Extract topics from markdown-style headers"""
        topics = {}
        
        # Pattern for headers (# Header, ## Header, etc.)
        header_pattern = re.compile(r'^(#{1,6})\s+(.+)$', re.MULTILINE)
        
        headers = []
        for match in header_pattern.finditer(text):
            level = len(match.group(1))
            title = match.group(2).strip()
            headers.append((level, title))
        
        if not headers:
            return topics
        
        # Build topic hierarchy from headers
        stack = []
        for level, title in headers:
            # Adjust stack based on level
            while stack and stack[-1][0] >= level:
                stack.pop()
            
            parent = stack[-1][1] if stack else None
            
            topic = Topic(
                name=title,
                level=level,
                parent=parent
            )
            topics[title] = topic
            
            stack.append((level, title))
        
        return topics
    
    async def _extract_topics_from_bullets(
        self,
        text: str,
        data: Dict[str, Any]
    ) -> Dict[str, Topic]:
        """Extract topics from bullet points"""
        topics = {}
        
        # Split into lines and analyze indentation
        lines = text.split('\n')
        bullet_pattern = re.compile(r'^(\s*)[•\-*]\s+(.+)$')
        
        current_parent = None
        indent_levels = {}
        
        for line in lines:
            match = bullet_pattern.match(line)
            if match:
                indent = len(match.group(1))
                content = match.group(2).strip()
                
                # Determine level based on indentation
                if indent not in indent_levels:
                    indent_levels[indent] = len(indent_levels)
                level = indent_levels[indent]
                
                # Find parent based on indentation
                if level > 0 and indent_levels:
                    parent_indent = min([k for k in indent_levels if indent_levels[k] == level - 1])
                    parent = current_parent if level > 1 else None
                else:
                    parent = None
                
                topic = Topic(
                    name=content,
                    level=level,
                    parent=parent
                )
                topics[content] = topic
                current_parent = content
        
        return topics
    
    async def _extract_topics_from_numbering(
        self,
        text: str,
        data: Dict[str, Any]
    ) -> Dict[str, Topic]:
        """Extract topics from numbered lists"""
        topics = {}
        
        # Pattern for numbered items (1.2.3, etc.)
        numbered_pattern = re.compile(
            r'^(\s*)((?:\d+\.)*\d+\.?)\s+(.+)$',
            re.MULTILINE
        )
        
        for match in numbered_pattern.finditer(text):
            indent = len(match.group(1))
            number = match.group(2)
            content = match.group(3).strip()
            
            # Determine level from number of dots
            level = number.count('.')
            
            # Find parent based on number
            if '.' in number:
                parent_num = number.rsplit('.', 1)[0]
                parent = f"{parent_num} {content}"  # Approximate
            else:
                parent = None
            
            topic = Topic(
                name=f"{number} {content}",
                level=level,
                parent=parent
            )
            topics[f"{number} {content}"] = topic
        
        return topics
    
    async def _extract_topics_from_toc(
        self,
        text: str,
        data: Dict[str, Any]
    ) -> Dict[str, Topic]:
        """Extract topics from table of contents"""
        topics = {}
        
        # Look for TOC section
        toc_patterns = [
            r'contents?[:\s]+(.*?)(?:\n\n|\Z)',
            r'table\s+of\s+contents?[:\s]+(.*?)(?:\n\n|\Z)',
            r'syllabus[:\s]+(.*?)(?:\n\n|\Z)'
        ]
        
        toc_text = None
        for pattern in toc_patterns:
            match = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
            if match:
                toc_text = match.group(1)
                break
        
        if toc_text:
            # Parse TOC entries (usually with page numbers)
            lines = toc_text.split('\n')
            for line in lines:
                # Remove page numbers
                clean_line = re.sub(r'\s+\d+$', '', line).strip()
                if clean_line and not clean_line.startswith('...'):
                    # Determine level based on indentation
                    indent = len(line) - len(line.lstrip())
                    level = indent // 2
                    
                    topic = Topic(
                        name=clean_line,
                        level=level
                    )
                    topics[clean_line] = topic
        
        return topics
    
    async def _extract_topics_from_sentences(self, text: str) -> Dict[str, Topic]:
        """Fallback: extract potential topics from sentences"""
        topics = {}
        
        sentences = sent_tokenize(text)
        
        # Look for sentences that might define topics
        topic_indicators = [
            r'unit\s+\d+[:\s]+(.+)',
            r'chapter\s+\d+[:\s]+(.+)',
            r'module\s+\d+[:\s]+(.+)',
            r'section\s+\d+[:\s]+(.+)',
            r'topic\s+\d+[:\s]+(.+)'
        ]
        
        for sentence in sentences[:20]:  # Check first 20 sentences
            for pattern in topic_indicators:
                match = re.search(pattern, sentence, re.IGNORECASE)
                if match:
                    topic_name = match.group(1).strip()
                    topic = Topic(name=topic_name)
                    topics[topic_name] = topic
                    break
        
        return topics
    
    async def _extract_prerequisites(
        self,
        text: str,
        topics: Dict[str, Topic]
    ) -> Dict[str, List[str]]:
        """Extract prerequisite relationships"""
        prerequisites = defaultdict(list)
        
        # Look for prerequisite statements
        prereq_patterns = [
            r'prerequisite[s]?[:\s]+(.+?)(?:\.|$)',
            r'assumes\s+(?:a\s+)?knowledge\s+of\s+(.+?)(?:\.|$)',
            r'requires?\s+(?:prior\s+)?(?:knowledge\s+of\s+)?(.+?)(?:\.|$)',
            r'should\s+have\s+(?:studied\s+)?(.+?)(?:\.|$)',
            r'must\s+have\s+(?:completed\s+)?(.+?)(?:\.|$)'
        ]
        
        for pattern in prereq_patterns:
            matches = re.finditer(pattern, text, re.IGNORECASE)
            for match in matches:
                prereq_text = match.group(1)
                
                # Extract individual prerequisites
                for topic_name in topics.keys():
                    if topic_name.lower() in prereq_text.lower():
                        # Find which topic this prerequisite applies to
                        # (usually the following topic)
                        context = text[match.end():match.end() + 200]
                        for next_topic in topics.keys():
                            if next_topic.lower() in context.lower():
                                prerequisites[next_topic].append(topic_name)
                                break
        
        return dict(prerequisites)
    
    def _extract_learning_outcomes(self, text: str) -> List[str]:
        """Extract learning outcomes/objectives"""
        outcomes = []
        
        # Look for learning outcomes section
        outcome_sections = [
            r'learning\s+outcomes?[:\s]+(.+?)(?:\n\n|\Z)',
            r'learning\s+objectives?[:\s]+(.+?)(?:\n\n|\Z)',
            r'what\s+you\'?ll\s+learn[:\s]+(.+?)(?:\n\n|\Z)',
            r'course\s+objectives?[:\s]+(.+?)(?:\n\n|\Z)'
        ]
        
        for pattern in outcome_sections:
            match = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
            if match:
                section = match.group(1)
                
                # Extract bullet points
                bullets = re.findall(r'[•\-*]\s*(.+?)(?:\n|$)', section)
                if bullets:
                    outcomes.extend(bullets)
                else:
                    # Extract sentences
                    sentences = sent_tokenize(section)
                    outcomes.extend(sentences)
                break
        
        # If no section found, look for outcome indicators in sentences
        if not outcomes:
            outcome_indicators = [
                r'will\s+be\s+able\s+to\s+(.+)',
                r'understand\s+(.+)',
                r'explain\s+(.+)',
                r'analyze\s+(.+)',
                r'evaluate\s+(.+)',
                r'create\s+(.+)'
            ]
            
            sentences = sent_tokenize(text)
            for sentence in sentences:
                for indicator in outcome_indicators:
                    if re.search(indicator, sentence, re.IGNORECASE):
                        outcomes.append(sentence)
                        break
        
        return outcomes
    
    def _extract_assessment_structure(self, text: str) -> Dict[str, Any]:
        """Extract assessment structure"""
        assessment = {
            'types': [],
            'weightage': {},
            'schedule': []
        }
        
        # Look for assessment section
        assessment_patterns = [
            r'assessment\s+(?:structure|pattern|scheme)[:\s]+(.+?)(?:\n\n|\Z)',
            r'evaluation\s+(?:scheme|pattern)[:\s]+(.+?)(?:\n\n|\Z)',
            r'marking\s+(?:scheme|pattern)[:\s]+(.+?)(?:\n\n|\Z)',
            r'exam\s+(?:pattern|structure)[:\s]+(.+?)(?:\n\n|\Z)'
        ]
        
        for pattern in assessment_patterns:
            match = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
            if match:
                section = match.group(1)
                
                # Extract assessment types and weightage
                weightage_pattern = re.compile(
                    r'([A-Za-z\s]+?)[:\s]+(\d+)(?:\s*%|\s*percent)',
                    re.IGNORECASE
                )
                for w_match in weightage_pattern.finditer(section):
                    assessment_type = w_match.group(1).strip()
                    weight = int(w_match.group(2))
                    assessment['types'].append(assessment_type)
                    assessment['weightage'][assessment_type] = weight
                
                break
        
        return assessment
    
    def _extract_references(self, text: str) -> Tuple[List[str], List[str]]:
        """Extract textbooks and reference materials"""
        textbooks = []
        references = []
        
        # Look for textbooks section
        textbook_patterns = [
            r'textbooks?[:\s]+(.+?)(?:\n\n|\Z)',
            r'prescribed\s+books?[:\s]+(.+?)(?:\n\n|\Z)',
            r'recommended\s+books?[:\s]+(.+?)(?:\n\n|\Z)'
        ]
        
        for pattern in textbook_patterns:
            match = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
            if match:
                section = match.group(1)
                # Extract book titles (usually with authors)
                books = re.findall(r'[•\-*\d+\.]\s*(.+?)(?:\n|$)', section)
                if books:
                    textbooks.extend(books)
                break
        
        # Look for references section
        ref_patterns = [
            r'references?[:\s]+(.+?)(?:\n\n|\Z)',
            r'further\s+reading[:\s]+(.+?)(?:\n\n|\Z)',
            r'bibliography[:\s]+(.+?)(?:\n\n|\Z)'
        ]
        
        for pattern in ref_patterns:
            match = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
            if match:
                section = match.group(1)
                refs = re.findall(r'[•\-*\d+\.]\s*(.+?)(?:\n|$)', section)
                if refs:
                    references.extend(refs)
                break
        
        return textbooks, references
    
    def _extract_exam_pattern(self, text: str) -> Dict[str, Any]:
        """Extract exam pattern details"""
        exam_pattern = {
            'duration': None,
            'total_marks': None,
            'sections': [],
            'question_types': [],
            'passing_marks': None
        }
        
        # Extract duration
        duration_match = re.search(
            r'duration[:\s]+(\d+)\s*(?:hours?|hrs?)',
            text,
            re.IGNORECASE
        )
        if duration_match:
            exam_pattern['duration'] = int(duration_match.group(1))
        
        # Extract total marks
        marks_match = re.search(
            r'total\s*marks?[:\s]+(\d+)',
            text,
            re.IGNORECASE
        )
        if marks_match:
            exam_pattern['total_marks'] = int(marks_match.group(1))
        
        # Extract question types
        question_types = [
            'multiple choice',
            'short answer',
            'long answer',
            'essay',
            'practical',
            'viva voce',
            'objective',
            'subjective'
        ]
        
        for q_type in question_types:
            if q_type in text.lower():
                exam_pattern['question_types'].append(q_type)
        
        return exam_pattern
    
    def _extract_title(self, text: str, data: Dict[str, Any]) -> str:
        """Extract syllabus title"""
        # Try from data first
        if isinstance(data, dict):
            for key in ['title', 'name', 'course', 'subject']:
                if key in data:
                    return str(data[key])
        
        # Try from text
        lines = text.strip().split('\n')
        if lines:
            # First line might be title
            first_line = lines[0].strip()
            if len(first_line) < 100:  # Reasonable title length
                return first_line
        
        return "Untitled Syllabus"
    
    def _extract_subject(self, text: str, data: Dict[str, Any]) -> str:
        """Extract subject name"""
        if isinstance(data, dict):
            for key in ['subject', 'course', 'discipline', 'field']:
                if key in data:
                    return str(data[key])
        
        # Common subject names
        subjects = [
            'mathematics', 'physics', 'chemistry', 'biology',
            'history', 'geography', 'economics', 'commerce',
            'computer science', 'english', 'hindi', 'sanskrit'
        ]
        
        text_lower = text.lower()
        for subject in subjects:
            if subject in text_lower:
                return subject.title()
        
        return "General"
    
    def _extract_grade_level(self, text: str, data: Dict[str, Any]) -> str:
        """Extract grade/class level"""
        if isinstance(data, dict):
            for key in ['grade', 'class', 'level', 'standard']:
                if key in data:
                    return str(data[key])
        
        # Try patterns
        for pattern, template in self.grade_patterns.items():
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                return template.format(match.group(1))
        
        return "Not Specified"
    
    def _extract_board(self, text: str, data: Dict[str, Any]) -> str:
        """Extract educational board"""
        if isinstance(data, dict) and 'board' in data:
            board_code = data['board'].lower()
            return self.boards.get(board_code, data['board'])
        
        text_lower = text.lower()
        for code, name in self.boards.items():
            if code in text_lower or name.lower() in text_lower:
                return name
        
        return "Not Specified"
    
    async def _enhance_syllabus(self, syllabus: Syllabus) -> Syllabus:
        """Enhance syllabus with additional information"""
        
        # Add learning objectives for each topic
        for topic_name, topic in syllabus.topics.items():
            # Extract key concepts
            topic.key_concepts = await self._extract_key_concepts(topic_name)
            
            # Generate learning objectives if missing
            if not topic.learning_objectives:
                topic.learning_objectives = await self._generate_learning_objectives(topic_name)
        
        # Build prerequisite graph
        syllabus.prerequisites = await self._build_prerequisite_graph(syllabus)
        
        return syllabus
    
    async def _extract_key_concepts(self, topic_name: str) -> List[str]:
        """Extract key concepts for a topic"""
        # This would ideally use NLP or knowledge base
        # For now, return topic name split into concepts
        words = topic_name.split()
        concepts = []
        
        # Extract noun phrases (simplified)
        for i, word in enumerate(words):
            if word[0].isupper() and i < len(words) - 1:
                concepts.append(f"{word} {words[i+1]}")
            elif word[0].isupper():
                concepts.append(word)
        
        return concepts[:5]  # Limit to 5 concepts
    
    async def _generate_learning_objectives(self, topic_name: str) -> List[str]:
        """Generate learning objectives for a topic"""
        objectives = [
            f"Understand the fundamental concepts of {topic_name}",
            f"Analyze key principles related to {topic_name}",
            f"Apply {topic_name} concepts to solve problems",
            f"Evaluate different aspects of {topic_name}"
        ]
        return objectives
    
    async def _build_prerequisite_graph(self, syllabus: Syllabus) -> Dict[str, List[str]]:
        """Build prerequisite relationship graph"""
        graph = defaultdict(list)
        
        # Add explicit prerequisites
        for topic, prereqs in syllabus.prerequisites.items():
            graph[topic].extend(prereqs)
        
        # Infer prerequisites based on topic hierarchy
        for topic_name, topic in syllabus.topics.items():
            if topic.parent:
                # Parent is prerequisite
                graph[topic_name].append(topic.parent)
        
        return dict(graph)
    
    def _compile_topic_patterns(self) -> List[re.Pattern]:
        """Compile regex patterns for topic detection"""
        return [
            re.compile(r'^(?:unit|chapter|module|section|lesson)\s+(\d+(?:\.\d+)*)[:\s]+(.+)$', re.I),
            re.compile(r'^(\d+(?:\.\d+)*)[:\s]+(.+)$'),
            re.compile(r'^#{1,6}\s+(.+)$'),
            re.compile(r'^[•\-*]\s+(.+)$')
        ]
    
    def _compile_objective_patterns(self) -> List[re.Pattern]:
        """Compile regex patterns for learning objectives"""
        return [
            re.compile(r'^[•\-*]\s*(?:understand|explain|describe|analyze|evaluate|create)\s+.+$', re.I),
            re.compile(r'^(?:will|should|shall)\s+be\s+able\s+to\s+.+$', re.I),
            re.compile(r'^learning\s+(?:outcome|objective)s?:\s*.+$', re.I)
        ]
    
    def _compile_assessment_patterns(self) -> List[re.Pattern]:
        """Compile regex patterns for assessment detection"""
        return [
            re.compile(r'assessment\s+(?:structure|pattern)', re.I),
            re.compile(r'evaluation\s+(?:scheme|pattern)', re.I),
            re.compile(r'marking\s+scheme', re.I),
            re.compile(r'exam\s+pattern', re.I)
        ]